from openai import OpenAI
import streamlit as st
import os
import time
from markdownify import markdownify as md
from docx import Document
from bs4 import BeautifulSoup


# from dotenv import load_dotenv


if 'api_submit_btn' not in st.session_state:
    st.session_state.api_submit_btn = False

def api_submit_btn():
    st.session_state.api_submit_btn = True

# Taking the api input in side ba
with st.sidebar:
    # form container starts from here
    api_form = st.form(key='api_input_form')
    api_form.text_input(':RED[ENTER YOUR API KEY]',key='api_key',placeholder='Insert your api key here', type="password")
    # Submit button in form container
    api_submit_button = api_form.form_submit_button(label='Submit',on_click=api_submit_btn)

api = st.session_state.api_key
# load_dotenv()
# api_key = os.getenv('OPENAI_API_KEY')
st.header('Code Documentation Assistant', divider="blue")

def llm(promt):
        response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
            "role": "system",
            "content": """You are a code documentation assistant//\n which helps creating the document for the code for developers.// \n//You will create the document in Markdown format.
            You will provide any suggestion related to improve code with heading 'Suggestions'// If you see very absurd code, //don't return anything just return,// ```Please provide correct code```// Return the document in HTML format//"""
            },
            # {
            # "role": "user",
            # "content": "If you see very absurd code, //don't return anything just return,// ```Please provide correct code```"
            # },
            {
            "role": "user",
            "content": f"""Elaborate all the functions, loops, if/else, variables or anything else that might be important and what is its use which you are using in the code.
            code: ```{prompt}```"""
            }
        ],

        temperature=0.8,
        max_tokens=1000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
        )
    
        return response


# api_key = os.environ['OPENAI_API_KEY_MINE']

client = OpenAI(api_key = api)
# with open('doc.json') as file:
#     # Step 2: Parse the JSON data
#     data = json.load(file) something
if 'text' not in st.session_state:
    st.session_state.text = False

def text_input():
    st.session_state.text = True

if 'clicked' not in st.session_state:
    st.session_state.clicked = False

def click_button():
    st.session_state.clicked = True
    st.session_state.download_clicked = False # Since code won't run after the download btn clicked, that's why making it false

if 'download_clicked' not in st.session_state:
    st.session_state.download_clicked = False

def click_download_button():
    st.session_state.download_clicked = True


# Display the selected option
time.sleep(3)

# form container starts from here
form = st.form(key='my_form')
# text area in form container
prompt = form.text_area(label="Code input",placeholder="Write your code here",height=305)
# Submit button in form container

submit_button = form.form_submit_button(label='Submit',on_click=click_button)

st.cache_resource
def cached_result(result):
    markdown_output = md(result)
    # Create a new Word document
    doc = Document()

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(result, 'html.parser')

    # Extract and add text and formatting to the document
    for tag in soup.find_all(['h1', 'p']):
        if tag.name == 'h1':
            doc.add_heading(tag.text, level=1)
        elif tag.name == 'p':
            doc.add_paragraph(tag.text)

    # Save the document in memory (in a variable)
    docx_output = doc
    return markdown_output, docx_output


if not st.session_state.download_clicked:

    if st.session_state.api_submit_btn and len(st.session_state.api_key) > 0:

        if st.session_state.clicked and api is not None and len(prompt) > 1:
            # Generating response from llm
            response = llm(prompt)

            # showing the total tokens used in sidebar
            with st.sidebar:
                f"Total tokens:   {response.usage.total_tokens}"

            # main result from the llm
            result = response.choices[0].message.content

            st.success("Your Doument is Generated successfully")

            #printing the result
            markdown_output, docx_output = cached_result(result)
            with st.expander("See the result"):
                markdown_output
            
            download_btn_docx = st.download_button(
                label="Download as word",
                key= "docx",
                data=markdown_output,
                file_name='result.md',
                mime='text',
                on_click=click_download_button
            )

            st.info("If you don't like the result you can submit again to regenerate", icon="ℹ️")
            print(result)


        elif st.session_state.clicked and api is not None:
            st.warning(":red[Please Write Some Code]") #Change this later
    elif st.session_state.api_submit_btn:
        with st.sidebar:
            st.warning(":red[Insert the API KEY]")
