from docx import Document
from bs4 import BeautifulSoup

# HTML content (replace with your actual HTML)
html_content = '''
<!DOCTYPE html>
<html>
<head>
    <title>Sample HTML</title>
    <style>
        h1 { color: blue; }
        p { font-size: 14px; }
    </style>
</head>
<body>
    <h1>Hello, Markdown!</h1>
    <p>This is a sample paragraph.</p>
</body>
</html>
'''

# Create a new Word document
doc = Document()

# Parse the HTML content using BeautifulSoup
soup = BeautifulSoup(html_content, 'html.parser')

# Extract and add text and formatting to the document
for tag in soup.find_all(['h1', 'p']):
    if tag.name == 'h1':
        doc.add_heading(tag.text, level=1)
    elif tag.name == 'p':
        doc.add_paragraph(tag.text)

# Save the document in memory (in a variable)
docx_output = doc

# Print the first paragraph (for demonstration purposes)
print(docx_output.paragraphs[0].text)
